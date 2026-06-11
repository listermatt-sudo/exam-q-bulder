from fastapi import FastAPI
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader

import requests
from io import BytesIO
import json
import os


# ✅ FastAPI app
app = FastAPI()

# ✅ CORS (safe even if same-origin)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ✅ Supabase base URL
BASE_URL = "https://gdcwjpkgffqmatsmuqra.supabase.co/storage/v1/object/public/question-images"


# ✅ Load structure.json
def load_structure():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, "structure.json")

    with open(file_path, "r") as f:
        return json.load(f)


structure_cache = load_structure()


# ✅ Request model
class RequestData(BaseModel):
    entries: list
    filetype: str


# ✅ Build filename
def build_base_name(month_year, paper, q):
    parts = month_year.split()
    month = parts[0]
    year = parts[1][-2:]

    if month == "November":
        month = "Nov"

    return f"{month} {year} {paper}_Q{q}"


# ✅ Load all pages for a question
def get_question_images(month_year, paper, q):

    base_name = build_base_name(month_year, paper, q)
    images = []

    # ✅ Try single-page first
    single_url = f"{BASE_URL}/{base_name}.png"

    try:
        r = requests.get(single_url)
        if r.status_code == 200:
            images.append(BytesIO(r.content))
            return images
    except:
        pass

    # ✅ Otherwise try multi-page versions
    page = 1

    while True:
        filename = f"{base_name}_{page}.png"
        url = f"{BASE_URL}/{filename}"

        try:
            r = requests.get(url)
        except:
            break

        if r.status_code != 200:
            break

        images.append(BytesIO(r.content))
        page += 1

    return images


# ✅ Serve frontend
@app.get("/", response_class=HTMLResponse)
def serve_frontend():
    with open("index.html", "r") as f:
        return f.read()


# ✅ Structure endpoint
@app.get("/structure")
def get_structure():
    return structure_cache


# ✅ ✅ ✅ WORD EXPORT (FULL WIDTH + NARROW MARGINS)
def create_word(entries, filename):

    doc = Document()

    # ✅ Set margins (~0.75 inch)
    section = doc.sections[0]
    section.top_margin = 700000
    section.bottom_margin = 700000
    section.left_margin = 700000
    section.right_margin = 700000

    doc.add_heading("Skills Map Exam Practice", 0)

    for paper, q in entries:

        doc.add_heading(f"{paper} — Q{q}", level=1)

        parts = paper.split()
        month = parts[0]
        year = parts[1]
        paper_code = parts[2]

        images = get_question_images(f"{month} {year}", paper_code, q)

        for img in images:

            # ✅ Compute available width dynamically
            section = doc.sections[0]
            available_width = (
                section.page_width
                - section.left_margin
                - section.right_margin
            )

            doc.add_picture(img, width=available_width)

        doc.add_page_break()

    doc.save(filename)


# ✅ PDF export (already well tuned)
def create_pdf(entries, filename):

    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4
    y = height - 40

    for paper, q in entries:

        parts = paper.split()
        month = parts[0]
        year = parts[1]
        paper_code = parts[2]

        images = get_question_images(f"{month} {year}", paper_code, q)

        for img_data in images:

            img_obj = ImageReader(img_data)
            img_w, img_h = img_obj.getSize()

            scale = (width - 80) / img_w
            new_h = img_h * scale

            if y - new_h < 50:
                c.showPage()
                y = height - 40

            c.drawImage(
                img_obj,
                40,
                y - new_h,
                width=width - 80,
                height=new_h
            )

            y -= new_h + 20

    c.save()


# ✅ Generate endpoint
@app.post("/generate")
def generate(data: RequestData):

    if data.filetype == "word":
        filename = "worksheet.docx"
        create_word(data.entries, filename)
    else:
        filename = "worksheet.pdf"
        create_pdf(data.entries, filename)

    return FileResponse(filename, filename=filename)


# ✅ Health check
@app.get("/health")
def health():
    return {"status": "ok"}
